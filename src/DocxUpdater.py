import os
from pathlib import Path

from tkinter import messagebox

from docx import Document
from docx.shared import Inches

class DocxUpdater:
    """ 
    Copypastes images in order of creation from a 'Screenshots' folder onto a 'Master.docx' file, in a specified path.
    """
    def __init__(self, root_dir):
        """
        Initializes the class with the root directory containing the folders.
        """
        self.root_dir = Path(root_dir)

    def get_screenshot_folders(self):
        """
        Get all subfolders that contain a 'Screenshots' directory.
        """
        screenshot_folders = []
        for root, dirs, files in os.walk(self.root_dir):
            if 'Screenshots' in dirs and 'Master.docx' in files:
                screenshot_folders.append(root)
        return screenshot_folders

    def get_img_files(self, screenshots_folder):
        """
        Retrieves and sorts image files by creation time in the 'Screenshots' folder. The image files must be either .jpg or .png .
        """
        png_files = []
        screenshots_path = Path(screenshots_folder)/'Screenshots' 

        if screenshots_path.is_dir():

            for file in screenshots_path.iterdir():
                
                #  Can we put something here to rename the img files? This way, we may avoid the issue with long paths.
                if file.suffix.lower() in {'.png','.jpg'}:
                    png_path = file
                    try: 
                        creation_time = png_path.stat().st_birthtime
                    except AttributeError:
                        creation_time = png_path.stat().st_ctime
                    png_files.append((file, creation_time))
        
        # Sort by creation time
        png_files.sort(key=lambda x: x[1])
        return [file[0] for file in png_files]
    
    def docx_contains_images(self,docx_path):
        """
        Returns True if the .docx file contains any inline images.
        """
        try:
            doc = Document(docx_path)
            return len(doc.inline_shapes) > 0
        except Exception as e:
            print(f"Error reading {docx_path}: {e}")
            return False

    def insert_images_to_docx(self, docx_path, png_files):
        """
        Inserts the sorted image files into the given .docx file.
        """
        from PIL import Image

        doc = Document(docx_path)
        
        # Are width and height necessary??
        # Is the truncation needed here?
        for png_file in png_files:
            # Open the image
            img = Image.open(str(png_file))
            width, height = img.size
            # Insert image into the docx file
            doc.add_picture(str(png_file), width=Inches(5))
            doc.add_paragraph(f"\n")
        
        doc.save(docx_path)

    def process_folders(self):
        """
        Iterates over all folders and inserts the screenshots into the corresponding Master.docx files.
        """
        folders = self.get_screenshot_folders()
       
        try:
            for folder in folders:
                docx_path = Path(folder)/'Master.docx'

                if not self.docx_contains_images(docx_path):
                
                # Get the .png files sorted by creation time
                    png_files = self.get_img_files(folder)
    
                # Insert the .png files into the Example.docx file
                    self.insert_images_to_docx(docx_path, png_files)

            # Message to confirm the successful operation
            messagebox.showinfo(title="Success!", message="Smistamento riuscito! Controlla i file master!")
        except Exception as e:
            messagebox.showerror(title="Error", message= f"Something went wrong: {e}")
