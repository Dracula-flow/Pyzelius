from pathlib import Path

from tkinter import filedialog, messagebox

from .Functions import truncate_path

class SanityTree:
    """
    Creates a directory based on a .xlsx file.

    The directory will consist of a main dir called "Sanity" and a variable number of subdirectories 
    each named after the rows in a specific column.

    Each subdirectory will have a "Screenshots" folder and a .docx file named "Master", with some 
    preformatted text (mainly the title, same as the subfolder).
    """
    def __init__(self, path):

        self.path = path
        self.screen_path= None
        self.df = None
        self.doc_path= None

    def new_master_dir(self):
        """
        Creates a new Master directory.
        """
        import pandas as pd
        from docx import Document
        
        path_to_excel = filedialog.askopenfilename()
        
        self.df = pd.read_excel(path_to_excel)
        
        for rows in self.df["Titolo"]:
            elaborated = truncate_path(rows)
            base = Path(self.path)/"Sanity"/elaborated
            self.screen_path = base /"Screenshots"
            self.screen_path.mkdir(parents=True,exist_ok=True)
            self.doc_path = base /"Master.docx"
            document = Document()
            document.add_heading(rows, 2)
            document.add_paragraph('BT=')
            document.save(self.doc_path)
        
        messagebox.showinfo(title="Cartella creata!", message="Cartella Sanity creata!")