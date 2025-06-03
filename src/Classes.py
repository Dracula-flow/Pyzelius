import os,json,sys,re
from pathlib import Path

import pandas as pd
import openpyxl as op
from openpyxl.worksheet.table import Table, TableStyleInfo
from docx import Document
from docx.shared import Inches
from PIL import Image
from tkinter import filedialog,messagebox

from src.Functions import time_responser, truncate_path

class CSV_File:
    """
    Creates two dataframes with the evidence data from the "Passed" and "Defects" folders in the Worktree. 
    This class works together with the 'Report' class.
    """
    def __init__(self, prefix:str, path:Path):
        self.prefix = prefix
        self.path = path
    
    @classmethod
    def passed_headers(cls):
        headers = ["Slot", "OS", "Clone","Test","Retest"]
        return headers

    @classmethod
    def defect_headers(cls):
        headers = ["Slot", "OS", "Clone","Test","Defect ID"]
        return headers
    
    def create_row(self,path:Path):
        """
        Creates an array of data based on the files in the directory passed to the method.
        """
        final_path = path/self.prefix 
        data = [f for f in final_path.iterdir() if f.is_file() and f.suffix in {'.mp4','.mov','.zip'}]
        rows = [re.split(r"-+",f.stem) for f in data]
        return rows
    
    def to_dataframe(self, path:Path ,headers: list[str])-> pd.DataFrame:
        """
        Passes the data into a dataframe
        """
        rows = self.create_row(path)
        return pd.DataFrame(rows, columns=headers)

# ---------------------------------------------------------------------------------------------------------------

class Report:
    """
    Creates an .xlsx file, printing two Pandas Dataframes on two different sheets. 
    The data in the dataframes must be supplied via the CSV_File class.
    """
    def __init__(self,path: Path):
        self.date_str = time_responser('date')
        self.path = path
        self.filename = self.path/f'Report_{self.date_str}.xlsx'


    def data_feed(self, df_passed:pd.DataFrame, df_defect:pd.DataFrame):
        """
        Creates the .xlsx file and loads the data from the CSV into it.
        """
        try:
            
            # CReates the xlsx file and writes the Passed data 
            df_passed.to_excel(self.filename, index=False, sheet_name="Passed")
            
            with pd.ExcelWriter(self.filename, engine="openpyxl", mode="a") as writer:

                 # Writes the Defects data
                df_defect.to_excel(writer, index=False, sheet_name="Defects")
            
            # Loads the sheets to an Openpyxl workbook
            wb = op.load_workbook(self.filename)

            for sheet_name in ["Passed","Defects"]:
            # Defines the two sheets inside the file and formats the columns
                ws = wb[sheet_name]
                self.adjust_column_width(ws)
                self.convert_to_table(ws)

            # Saves the file
            wb.save(self.filename)
            messagebox.showinfo(title="Report creato!", message=f"Il report {self.date_str} è stato creato con successo!")

        except pd.errors.ParserError:
            messagebox.showerror(title="Error!", message="A file does not respect the format!")

    def adjust_column_width(self, worksheet):
        """
        Adapts the width of the columns based on the content.
        """
        for col in worksheet.columns:
            try:
                max_length = 0
                column = col[0].column_letter 
                for cell in col:
                        if len(cell.value) > max_length:
                            max_length = max(max_length,len(cell.value))
                adjusted_width = (max_length + 2) 
                worksheet.column_dimensions[column].width = adjusted_width
                # It's gonna raise a TypeError exception anytime there's an empty field in the CSV
                # That's gonna happen sometimes, so it's pass.
            except TypeError:
                pass
    
    def convert_to_table(self, worksheet):
        """
        Converts the used range of the worksheet into an Excel table with styling.
        """
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        if max_row < 2:
            return  # No data to table-ify

        ref = f"A1:{worksheet.cell(row=max_row, column=max_col).coordinate}"
        table = Table(displayName=f"Table_{worksheet.title}", ref=ref)

        # Apply a table style
        style = TableStyleInfo(
            name="TableStyleMedium2", showFirstColumn=False,
            showLastColumn=False, showRowStripes=True, showColumnStripes=False
        )
        table.tableStyleInfo = style

        worksheet.add_table(table)
# ------------------------------------------------------------------------------------------------------------------------

class WorkTree:
    """
    Creates a directory named after the day, with three sub-directories named Passed, Defects and Report.
    Passed and Defects will contain the video test-evidence, Report will host the Report generated via the "Report" class.
    """
    def __init__(self,path):
        self.path = path
        self.dirname = time_responser('date')
        self.subdirs = ("Passed","Defects", "Report")

    def create_worktree(self):
        root_path = Path(self.path)/self.dirname
        try:
            Path.mkdir(root_path)
            messagebox.showinfo(title="Cartella creata!", message=f"Cartella {self.dirname} creata!")
        except FileExistsError:
            messagebox.showerror(title="Errore!",message=f"La cartella {self.dirname} esiste già nella destinazione!")

        for subdir in self.subdirs:
            subdir_path = root_path/subdir
            Path.mkdir(subdir_path)

# -------------------------------------------------------------------------------------------------------------------

class Pathfinder:
    """
    Registers the user's preferred destination for the test evidence and allows its modification. 

    The data will be registered into a .JSON file.
    
    While in source code mode, it will default to the "config/config.json" file present in the root.
    
    While in executable mode, if absent, it will create and default to a "pyzelius_config" folder in the Home directory.

    If the path is modified, the app will need to be restarted.

    """
    def __init__(self):
        self.config_last_path = self.get_config_path()
        self.last_path = self.load_last_path()

    def get_config_path(self):
        """
        Returns the path to the configuration file, ensuring it's writable.
        """
        if hasattr(sys, '_MEIPASS'):
            # If running as a bundled executable, save it to the user's home directory or app-specific folder
            config_folder = Path.home()/'pyzelius_config'
            if not Path.exists(config_folder):
                config_folder.mkdir(parents=True, exist_ok=False)
            config_path = config_folder/"config.json"

            if not config_path.exists():
                default_config = {}  # or provide a meaningful default structure
                with config_path.open('w', encoding='utf-8') as f:
                    json.dump(default_config, f, indent=4)
            return config_path
            
        else:
            # If running as a Python script, use the relative path
            return Path(__file__).resolve().parent.parent / 'config/config.json'
        
    def load_last_path(self):
        """
        Load the last used directory path from a configuration file.
        """
        if Path.exists(self.config_last_path):
            try:
                with open(self.config_last_path, 'r') as config_file:
                    config = json.load(config_file)
                    return config.get("last_path", "")
            except PermissionError:
                print ("Permission error while trying to read the file")
            except json.JSONDecodeError:
                print("Error decoding the configuration file.")
                return ""
        return ""  # Return an empty string if no path was saved

    def save_last_path(self):
        """
        Save the current directory path to the configuration file.
        """
        config = {"last_path": self.last_path}
        try:
            with open(self.config_last_path, 'w') as config_file:
                json.dump(config, config_file, indent=4)
        except PermissionError:
            messagebox.showerror(title="Error!", message="Permission error in attempt to write to the configuration file.")

    def modify_last_path(self):
        """
        Allows the user to modify the preferred path.
        """
        
        new_path = filedialog.askdirectory()

        if new_path:
            self.last_path = new_path
            self.save_last_path()

    def get_path(self):
        return self.last_path
    
# --------------------------------------------------------------------------------------------------------------------

class Signature():
    """
    Generates a signature necessary to close a test. It has to be mounted on a Tkinter frame.
    """
    def __init__(self):
        self.input_fields = ("Sigla", 
                "BT",
                "APP",
                "BUILD", 
                "DEVICE",
                "OS",
                "Puntamento")
        self.combine=[]

    def entry_combine(self, entry_list):
        self.combine.clear()
        for entry,key in zip(entry_list,self.input_fields):
            try:
                # input_value= entry.get()
                self.combine.append(f"{key} {entry}")
            except TypeError:
                messagebox.showerror (f"Error: Invalid Input - {key}")

        self.combine.append(f"{'DATA E ORA'} {time_responser('datetime')}")
        self.result = "\n".join(self.combine)
        return self.result
    
class SignatureSanity(Signature):
    """
    Signature subclass with fewer fields.
    """
    def __init__(self):
        super().__init__()
        self.input_fields = ("Sigla", "BT", "CLONE", "BROWSER",)  # Subclass with fewer fields

class SignatureMinimal(Signature):
    """
    Signature class with just the initials.
    """
    def __init__(self):
        super().__init__()
        self.input_fields = ("Sigla",) # The comma is necessary, otherwise the entry_combine will interpret this as a string to split
# -----------------------------------------------------------------------------------------------------------------------------------------

class Master:
    """
    Creates a directory based on a .xlsx file.

    The directory will consist of a main dir called "Sanity" and a variable number of subdirectories 
    each named after the rows in a specific column.

    Each subdirectory will have a "Screenshots" folder and a .docx file named "Master", with some 
    preformatted text (mainly the title, same as the subfolder).
    """
    def __init__(self, path):

        self.path = path
        self.screen_path= None
        self.df = None
        self.doc_path= None

    def new_master_dir(self):
        """
        Creates a new Master directory.
        """
        path_to_excel = filedialog.askopenfilename()
        
        self.df = pd.read_excel(path_to_excel)
        
        for rows in self.df["Titolo"]:
            elaborated = truncate_path(rows)
            base = Path(self.path)/"Sanity"/elaborated
            self.screen_path = base /"Screenshots"
            self.screen_path.mkdir(parents=True,exist_ok=True)
            self.doc_path = base /"Master.docx"
            document = Document()
            document.add_heading(rows, 2)
            document.add_paragraph('BT=')
            document.save(self.doc_path)
        
        messagebox.showinfo(title="Cartella creata!", message="Cartella Sanity creata!")

# ------------------------------------------------------------------------------------------------------
# Needed: Error Handling, also trying to enable long paths on the folders.
# Enabling long path on the machine is not the solution. Truncating the path is seemingly doing nothing.

class DocxUpdater:
    """ 
    Copypastes images in order of creation from a 'Screenshots' folder onto a 'Master.docx' file, in a specified path.
    """
    def __init__(self, root_dir):
        """
        Initializes the class with the root directory containing the folders.
        """
        self.root_dir = Path(root_dir)

    def get_screenshot_folders(self):
        """
        Get all subfolders that contain a 'Screenshots' directory.
        """
        screenshot_folders = []
        for root, dirs, files in os.walk(self.root_dir):
            if 'Screenshots' in dirs and 'Master.docx' in files:
                screenshot_folders.append(root)
        return screenshot_folders

    def get_img_files(self, screenshots_folder):
        """
        Retrieves and sorts image files by creation time in the 'Screenshots' folder. The image files must be either .jpg or .png .
        """
        png_files = []
        screenshots_path = Path(screenshots_folder)/'Screenshots' 

        if screenshots_path.is_dir():

            for file in screenshots_path.iterdir():
                
                #  Can we put something here to rename the img files? This way, we may avoid the issue with long paths.
                if file.suffix.lower() in {'.png','.jpg'}:
                    png_path = file
                    try: 
                        creation_time = png_path.stat().st_birthtime
                    except AttributeError:
                        creation_time = png_path.stat().st_ctime
                    png_files.append((file, creation_time))
        
        # Sort by creation time
        png_files.sort(key=lambda x: x[1])
        return [file[0] for file in png_files]
    
    def docx_contains_images(self,docx_path):
        """
        Returns True if the .docx file contains any inline images.
        """
        try:
            doc = Document(docx_path)
            return len(doc.inline_shapes) > 0
        except Exception as e:
            print(f"Error reading {docx_path}: {e}")
            return False

    def insert_images_to_docx(self, docx_path, png_files):
        """
        Inserts the sorted image files into the given .docx file.
        """
        doc = Document(docx_path)
        
        # Are width and height necessary??
        # Is the truncation needed here?
        for png_file in png_files:
            # Open the image
            img = Image.open(str(png_file))
            width, height = img.size
            # Insert image into the docx file
            doc.add_picture(str(png_file), width=Inches(5))
            doc.add_paragraph(f"\n")
        
        doc.save(docx_path)

    def process_folders(self):
        """
        Iterates over all folders and inserts the screenshots into the corresponding Master.docx files.
        """
        folders = self.get_screenshot_folders()
       
        try:
            for folder in folders:
                docx_path = Path(folder)/'Master.docx'

                if not self.docx_contains_images(docx_path):
                
                # Get the .png files sorted by creation time
                    png_files = self.get_img_files(folder)
    
                # Insert the .png files into the Example.docx file
                    self.insert_images_to_docx(docx_path, png_files)

            # Message to confirm the successful operation
            messagebox.showinfo(title="Success!", message="Smistamento riuscito! Controlla i file master!")
        except Exception as e:
            messagebox.showerror(title="Error", message= f"Something went wrong: {e}")

#----------------------------------------------------------------------------------------------------------------------------

class DeviceUpdater:
    """
    Handles the updating of the Devices memorized in the JSON file.
    """
    def __init__(self, json_path: Path):
        self.json_path = json_path
        self.config_data = self.device_getter()

    def device_getter(self):
        with open(self.json_path, "r") as f:
            return json.load(f)
        
    def ensure_device_key(self):
        if "Devices" not in self.config_data:
            self.config_data["Devices"] = {}  
    
    def JSON_writer(self):
        with open(self.json_path, "w") as w:
            json.dump(self.config_data, w, indent=4)

    def add_entry(self, new_device: str, new_os: str):
        "Adds a device entry to the JSON file"
        # Ensure that the "Devices" key exists in the JSON structure
        self.ensure_device_key()

        # Generate a new ID based on the existing ones
        new_id = int(len(self.config_data["Devices"]) + 1)

        # Add the new device entry with the generated ID
        self.config_data["Devices"][new_id] = {
            "device": new_device,
            "os": new_os
        }

        # Write the updated JSON data back to the file
        self.JSON_writer()

    def modify_entry(self, device_id: int, new_device: str, new_os: str):
        "Modifies a device entry by ID, provided there are new arguments"
        self.ensure_device_key()

             # Check if the device ID exists
        if device_id in self.config_data["Devices"]:
            # Modify the device entry
            self.config_data["Devices"][device_id]["device"] = new_device
            self.config_data["Devices"][device_id]["os"] = new_os
            self.JSON_writer()  # Save changes to the file
        else:
            raise ValueError(f"Device ID '{device_id}' does not exist in the configuration.")

        self.JSON_writer()

    def delete_entry(self, device_id:str):
        """Deletes a device entry by ID."""
        self.ensure_device_key()

        # Check if the device ID exists
        if device_id in self.config_data["Devices"]:
            # Remove the device entry
            del self.config_data["Devices"][device_id]
            self.JSON_writer()  # Save changes to the file
        else:
            raise ValueError(f"Device ID '{device_id}' does not exist in the configuration.")
    
    def load_devices(self, tree: object):
        try:
            with open(self.json_path, "r") as file:
                data = json.load(file)
                devices = data.get("Devices", {})
                for device_id, device_info in devices.items():
                    tree.insert("", "end", values=(device_id, device_info["device"], device_info["os"]))
        except (FileNotFoundError, json.JSONDecodeError) as e:
            print(f"Error loading devices: {e}")

